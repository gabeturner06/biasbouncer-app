import tempfile
import aiofiles
import os
import streamlit as st
import fitz
import json
import csv
import pandas as pd
import docx
from docx import Document
import mimetypes
from openpyxl import Workbook

# Ensure session state has a temp directory
def ensure_temp_dir():
    if "temp_dir" not in st.session_state:
        st.session_state["temp_dir"] = tempfile.mkdtemp(prefix="biasbouncer_")
    return st.session_state["temp_dir"]

# Function to list files in the session's temp directory
def list_files():
    temp_dir = ensure_temp_dir()  # Ensure temp directory is initialized
    return os.listdir(temp_dir)

# Function to write a file and trigger UI update

async def write_tool(filename: str, content: str):
    try:
        temp_dir = ensure_temp_dir()  # Ensure temp directory is initialized
        temp_file_path = os.path.join(temp_dir, filename)

        # Determine file extension
        file_ext = filename.lower().split('.')[-1]

        if file_ext == "txt":
            async with aiofiles.open(temp_file_path, mode='w', encoding='utf-8') as file:
                await file.write(content)
        
        elif file_ext == "md":
            async with aiofiles.open(temp_file_path, mode='w', encoding='utf-8') as file:
                await file.write(content)

        elif file_ext == "py":
            async with aiofiles.open(temp_file_path, mode='w', encoding='utf-8') as file:
                await file.write(content)

        elif file_ext == "pdf":
            doc = fitz.open()
            page = doc.new_page()
            
            text = content.replace("\n", " ")  # Ensure line breaks are handled properly
            text_rect = fitz.Rect(50, 50, 550, 800)  # Define text area on the page
            
            page.insert_textbox(text_rect, text, fontsize=12, fontname="helv", align=0)
            doc.save(temp_file_path)


        elif file_ext == "docx":
            doc = Document()
            for paragraph in content.split("\n"):  # Ensure paragraphs are separated properly
                doc.add_paragraph(paragraph)
            doc.save(temp_file_path)


        elif file_ext == "xlsx":
            wb = Workbook()
            ws = wb.active
            for i, line in enumerate(content.split("\n"), start=1):
                cells = line.split("\t") if "\t" in line else line.split(",")  # Handle CSV or tab-separated data
                for j, cell in enumerate(cells, start=1):
                    ws.cell(row=i, column=j, value=cell)
            wb.save(temp_file_path)

        elif file_ext == "csv":
            with open(temp_file_path, mode="w", newline="", encoding="utf-8") as file:
                writer = csv.writer(file)
                for line in content.split("\n"):
                    writer.writerow(line.split(",")) 


        else:
            return f"Error: Unsupported file type '{file_ext}'. Supported formats: TXT, PDF, DOCX, XLSX."

        st.session_state["file_updated"] = True  # Trigger UI refresh
        return f"✅ Successfully wrote to '{temp_file_path}'."

    except Exception as e:
        return f"❌ Error writing to file: {str(e)}"

async def read_tool(filename: str):
    try:
        temp_dir = ensure_temp_dir()
        temp_file_path = os.path.join(temp_dir, filename)

        if not os.path.exists(temp_file_path):
            return f"Error: Temporary file '{filename}' does not exist."

        # Determine file extension
        file_ext = filename.lower().split('.')[-1]

        if file_ext == "txt":
            async with aiofiles.open(temp_file_path, mode='r', encoding='utf-8') as file:
                await file.read(content)
        
        elif file_ext == "md":
            async with aiofiles.open(temp_file_path, mode='r', encoding='utf-8') as file:
                await file.read(content)

        elif file_ext == "py":
            async with aiofiles.open(temp_file_path, mode='r', encoding='utf-8') as file:
                await file.read(content)

        elif file_ext == "pdf":
            content = await read_pdf(temp_file_path)

        elif file_ext == "csv":
            content = await read_csv(temp_file_path)

        elif file_ext == "json":
            content = await read_json(temp_file_path)

        elif file_ext in ["xls", "xlsx"]:
            content = await read_excel(temp_file_path)

        elif file_ext == "docx":
            content = await read_docx(temp_file_path)

        else:
            mime_type, _ = mimetypes.guess_type(temp_file_path)
            return f"Error: Unsupported file type '{file_ext}' (MIME type: {mime_type})."

        return content

    except Exception as e:
        return f"Error reading from file: {str(e)}"

# PDF Reader
async def read_pdf(pdf_path: str):
    """Extracts text from a PDF file asynchronously."""
    try:
        doc = fitz.open(pdf_path)
        text = "\n".join([page.get_text("text") for page in doc])
        return text if text else "Warning: No text found in the PDF."
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

# CSV Reader
async def read_csv(csv_path: str):
    """Reads a CSV file and returns its content as a string."""
    try:
        with open(csv_path, newline='', encoding='utf-8') as file:
            reader = csv.reader(file)
            content = "\n".join([", ".join(row) for row in reader])
        return content
    except Exception as e:
        return f"Error reading CSV file: {str(e)}"

# JSON Reader
async def read_json(json_path: str):
    """Reads a JSON file and returns its content as a formatted string."""
    try:
        with open(json_path, "r", encoding="utf-8") as file:
            data = json.load(file)
        return json.dumps(data, indent=4)  # Pretty print the JSON content
    except Exception as e:
        return f"Error reading JSON file: {str(e)}"

# Excel Reader
async def read_excel(excel_path: str):
    """Reads an Excel file and returns the first few rows as a string."""
    try:
        df = pd.read_excel(excel_path)
        return df.head().to_string(index=False)  # Convert the first few rows to a string
    except Exception as e:
        return f"Error reading Excel file: {str(e)}"

# DOCX Reader
async def read_docx(docx_path: str):
    """Reads a Word document and extracts its text."""
    try:
        doc = docx.Document(docx_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text if text else "Warning: No text found in the DOCX file."
    except Exception as e:
        return f"Error reading DOCX file: {str(e)}"
